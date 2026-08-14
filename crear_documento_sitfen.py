from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(__file__).parent
OUT = ROOT / "Documento_Arquitectura_SITFEN.docx"

NAVY = "102A43"
BLUE = "1769AA"
TEAL = "087F8C"
PURPLE = "553C9A"
ORANGE = "C05621"
GREEN = "2F855A"
LIGHT = "F3F7FA"
MID = "D9E2EC"
TEXT = "243B53"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=MID, sz="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color=TEXT, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, header_color=NAVY, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True, "FFFFFF", font_size)
        shade(table.rows[0].cells[i], header_color)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), False, TEXT, font_size)
            shade(cells[i], "FFFFFF" if ridx % 2 == 0 else LIGHT)
            set_cell_border(cells[i])
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_border(cell)
            if widths and i < len(widths):
                cell.width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.55 + level * 0.35)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_note(doc, title, text, fill="EAF4F4", border=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = table.cell(0, 0)
    shade(c, fill)
    set_cell_border(c, border, "10")
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(border)
    r.font.size = Pt(9.5)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor.from_string(TEXT)
    r2.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("SITFEN · Arquitectura de solución · ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("829AB1")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.65)
sec.bottom_margin = Cm(1.55)
sec.left_margin = Cm(1.75)
sec.right_margin = Cm(1.75)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9.5)
styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT)
styles["Normal"].paragraph_format.space_after = Pt(5)
for name, size, color in [("Heading 1", 17, NAVY), ("Heading 2", 12.5, BLUE), ("Heading 3", 10.5, TEAL)]:
    st = styles[name]
    st.font.name = "Aptos Display"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(10)
    st.paragraph_format.space_after = Pt(5)

header = sec.header.paragraphs[0]
header.text = "SITFEN  |  Smart Agentic Chatbot v2"
header.runs[0].font.size = Pt(8)
header.runs[0].font.color.rgb = RGBColor.from_string("829AB1")
add_page_number(sec.footer.paragraphs[0])

# Portada
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(22)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SITFEN")
r.font.name = "Aptos Display"
r.font.size = Pt(34)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(NAVY)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("Arquitectura de solución e infraestructura de implementación agéntica")
r.font.size = Pt(18)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(TEAL)
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p3.add_run("Asistente conversacional para orientar a empresarios agroindustriales ante el Fenómeno El Niño")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor.from_string(TEXT)
doc.add_paragraph()
cover = doc.add_table(rows=1, cols=1)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
c = cover.cell(0, 0)
shade(c, "EAF4F4")
set_cell_border(c, TEAL, "12")
c.text = ""
pc = c.paragraphs[0]
pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pc.add_run("Documento técnico\n")
rr.bold = True
rr.font.size = Pt(12)
rr.font.color.rgb = RGBColor.from_string(TEAL)
rr2 = pc.add_run("Perfil de agente de referencia: Agent Profile Card – SITFEN v2\n\n")
rr2.font.size = Pt(10)
rr2.font.color.rgb = RGBColor.from_string(TEXT)
rr3 = pc.add_run("Versión de arquitectura: 1.0  |  06 de agosto de 2026")
rr3.font.size = Pt(9)
rr3.font.color.rgb = RGBColor.from_string(TEXT)
doc.add_page_break()

# Resumen
add_heading(doc, "1. Resumen ejecutivo")
doc.add_paragraph(
    "SITFEN se plantea como un asistente agéntico semi-autónomo y constreñido: identifica el punto de dolor del empresario o productor agroindustrial, clasifica la necesidad y recomienda la entidad pública peruana más idónea. El agente no otorga subsidios ni ejecuta trámites en nombre del Estado; su responsabilidad es orientar, entregar canales oficiales y escalar los casos complejos, urgentes o de alto impacto a una persona."
)
doc.add_paragraph(
    "La arquitectura propuesta combina un canal conversacional por WhatsApp, un núcleo de orquestación basado en LangGraph/LangChain y Gemini vía Vertex AI, recuperación aumentada por generación (RAG) sobre fuentes oficiales, memoria de corto y largo plazo, herramientas de derivación y un conjunto de controles de seguridad, trazabilidad y supervisión humana. La ejecución en Google Cloud Platform prioriza servicios serverless para escalar bajo demanda y controlar el costo del MVP."
)
add_note(doc, "Decisión arquitectónica central", "La autonomía se limita a un catálogo de entidades, programas y reglas de derivación validadas. El modelo interpreta y coordina; las reglas y las fuentes oficiales restringen qué puede recomendar.")

# Profile mapping
add_heading(doc, "2. Traducción del perfil del agente a requisitos de solución")
rows = [
    ("Dominio", "Empresarios y productores agroindustriales de Piura, Lambayeque, La Libertad y Tumbes; afectaciones y riesgos asociados al FEN."),
    ("Objetivo", "Detectar emoción y necesidad; clasificar información, soporte técnico, financiamiento, subsidio, insumos o materiales; derivar a la entidad competente."),
    ("Conocimiento", "Mapa institucional, fondos y programas, protocolos SINAGERD, requisitos y procedimientos vigentes; corpus administrable y versionado."),
    ("Interacción", "WhatsApp como canal principal; texto, voz e imagen; lenguaje empático y respuesta con enlaces/canales oficiales."),
    ("Memoria", "Estado de la sesión, ubicación, cultivo/actividad, necesidad; historial de derivaciones, seguimiento, feedback y derecho al olvido."),
    ("Autonomía", "Semi-autónomo, constreñido por catálogo y reglas de derivación; no realiza trámites ni promete beneficios."),
    ("Criticidad", "Medio-alto controlado: se requiere escalamiento humano, información vigente, trazabilidad, protección de datos y manejo empático."),
]
add_table(doc, ["Dimensión del perfil", "Implicación para la arquitectura"], rows, [3.6, 13.1], header_color=TEAL)

# Architecture image
add_heading(doc, "3. Gráfica de la arquitectura actualizada")
doc.add_paragraph("La siguiente vista integra los canales, las APIs de entrada, el núcleo agéntico, las herramientas, los datos, las integraciones externas y las salidas esperadas según el perfil SITFEN.")
arch = ROOT / "sitfen_arquitectura_blanco.png"
if arch.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(arch), width=Cm(17.0))
    cap = doc.add_paragraph("Figura 1. Arquitectura funcional de la solución SITFEN.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(8)
else:
    add_note(doc, "Recurso gráfico no encontrado", "La imagen de arquitectura no se encontró en el directorio del proyecto.", "FFF5EB", ORANGE)

add_heading(doc, "3.1 Flujo principal de atención")
add_table(doc, ["Paso", "Flujo", "Control o resultado"], [
    ("1", "El empresario envía texto, nota de voz o imagen por WhatsApp.", "Twilio recibe el mensaje y entrega el webhook al servicio agéntico."),
    ("2", "Se valida firma, se limita la tasa y se anonimiza el identificador.", "Se evita spoofing, abuso y exposición de PII en trazas o memoria."),
    ("3", "Se recuperan memoria corta, resumen histórico y contexto semántico.", "El agente entiende la situación sin depender del estado de un proceso."),
    ("4", "LangGraph ReAct selecciona herramientas: RAG, clima, derivación o handoff.", "El uso de tools está acotado por el prompt, guardrails y reglas deterministas."),
    ("5", "Gemini interpreta, redacta y decide la secuencia de acciones.", "Router Flash/Pro según complejidad y costo; safety settings activos."),
    ("6", "Se responde con orientación, entidad, canal oficial y fuentes.", "Respuesta auditable, sin falsas promesas; se registra la derivación."),
    ("7", "Si el caso es urgente, complejo o sensible, se escala a un operador.", "Ticket de handoff con contexto y consola en vivo para atención humana."),
], [1.1, 8.2, 7.4], header_color=BLUE, font_size=8.2)

# Infrastructure
add_heading(doc, "4. Infraestructura de la implementación agéntica")
doc.add_paragraph("La infraestructura se organiza como una plataforma serverless en GCP, con servicios stateless para la atención en línea y procesos batch para la ingesta, evaluación y automatización.")
infra_rows = [
    ("Canal e ingreso", "Twilio WhatsApp; webhook FastAPI; API web REST; APIs protegidas de consola; portal SST.", "Cloud Run + HTTPS; validación X-Twilio-Signature fail-closed."),
    ("Servicios de aplicación", "Servicio del agente, Admin UI, portal SST/consola, job de ingesta y job de evaluación.", "Cloud Run Services y Cloud Run Jobs, cada uno con imagen Docker independiente cuando corresponde."),
    ("Núcleo agéntico", "AgentOrchestrator, guardrails, router determinista Flash/Pro, agente ReAct y 10 tools.", "Python, FastAPI, LangGraph y LangChain."),
    ("Capa de IA", "Generación, clasificación, reranking, embeddings, transcripción, voz y visión.", "Gemini 2.5 Flash/Pro, Vertex AI, embeddings multilingües, Speech-to-Text/Text-to-Speech y Gemini multimodal."),
    ("Datos", "Memoria, usuarios, derivaciones, eventos, vectores, corpus documental, audios y assets.", "Firestore/Vector Search y Cloud Storage; cifrado en reposo e IAM."),
    ("Eventos y automatización", "Informe diario, re-engagement, alertas SST y tareas diferidas.", "Cloud Scheduler, Pub/Sub, Cloud Tasks y/o Eventarc según el flujo."),
    ("Seguridad y operación", "Secretos, identidades, logs, métricas, trazas, alertas y control de gasto.", "Secret Manager, IAM, Cloud Logging, Monitoring, Trace, Artifact Registry, Cloud Build y presupuestos."),
]
add_table(doc, ["Capa", "Responsabilidad", "Implementación propuesta"], infra_rows, [3.3, 7.0, 6.4], header_color=PURPLE, font_size=8.1)

gcp = ROOT / "sitfen_despliegue_gcp_blanco_preservado.png"
if gcp.exists():
    doc.add_paragraph("La distribución física de los componentes en GCP se resume en la Figura 2.")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(gcp), width=Cm(16.8))
    cap = doc.add_paragraph("Figura 2. Despliegue de SITFEN en Google Cloud Platform.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic = True; cap.runs[0].font.size = Pt(8)

add_heading(doc, "4.1 Ambientes y despliegue")
add_bullets(doc, [
    "Desarrollo: ejecución local con Docker y emuladores/fakes para Firestore cuando no se requiere conexión a GCP.",
    "Pruebas: validación unitaria, pruebas de integración del webhook y evaluación del RAG con dataset dorado; sin datos personales reales.",
    "Producción: despliegue de imágenes en Artifact Registry mediante Cloud Build/Cloud Deploy hacia Cloud Run, con secretos montados desde Secret Manager.",
    "Procesamiento batch: Cloud Run Jobs para ingestión de PDF/DOCX/TXT/URL, deduplicación por SHA-256, chunking, embeddings y evaluación RAGAS/DeepEval.",
    "Escalamiento: Cloud Run sin estado y escalado automático; Firestore y Storage como servicios administrados; límites de concurrencia, timeout y presupuesto configurados.",
])

# Inventory
add_heading(doc, "5. Herramientas, componentes, frameworks y servicios a desplegar")
stack_rows = [
    ("Python", "Lenguaje base", "Servicios agénticos, ingesta, APIs y evaluación."),
    ("FastAPI / Uvicorn", "Framework API", "Webhook WhatsApp, API web, voz, consola y health checks."),
    ("LangGraph", "Orquestación agéntica", "Agente ReAct, ciclo pensamiento–acción–observación y límite de recursión."),
    ("LangChain", "Abstracciones LLM/tools", "Prompts, herramientas, retriever y conexión con Vertex AI."),
    ("Gemini 2.5 Flash / Pro", "Modelos generativos", "Clasificación, respuesta, resumen, reranking y razonamiento según complejidad."),
    ("Vertex AI", "Plataforma de IA", "Inferencia, embeddings, multimodalidad y Speech-to-Text/Text-to-Speech."),
    ("Firestore + Vector Search", "Persistencia y recuperación", "Usuarios, conversaciones, memoria, handoffs, eventos y KNN vectorial."),
    ("Cloud Storage", "Objetos y documentos", "Corpus, audio, imágenes, assets del avatar y archivos de evaluación."),
    ("Twilio WhatsApp API", "Mensajería", "Canal principal; webhooks, plantillas y envío de respuesta."),
    ("Streamlit", "Interfaz administrativa", "Carga de conocimiento, control de usuarios y seguimiento de ingesta."),
    ("Cloud Run", "Compute serverless", "Agente, Admin UI, portal SST y consola web."),
    ("Cloud Run Jobs", "Procesamiento batch", "Ingesta reproducible y evaluación offline."),
    ("Cloud Scheduler / Pub/Sub / Tasks", "Orquestación", "Informes, re-engagement, alertas y trabajos asíncronos."),
    ("Secret Manager / IAM", "Seguridad", "Secretos, identidades, permisos mínimos y separación de responsabilidades."),
    ("Cloud Logging / Monitoring / Trace", "Observabilidad GCP", "Logs estructurados, métricas, alertas, trazabilidad y diagnóstico."),
    ("LangSmith / Confident AI", "Observabilidad y evaluación", "Trazas de ejecución y monitoreo/evaluación del agente en ambientes habilitados."),
    ("RAGAS / DeepEval", "Evaluación", "Faithfulness, relevancia, recuperación y calidad de derivación sobre dataset dorado."),
    ("Docker / Artifact Registry / Cloud Build", "DevSecOps", "Imágenes reproducibles, registro, build y promoción controlada."),
]
add_table(doc, ["Elemento", "Tipo", "Uso en SITFEN"], stack_rows, [4.1, 4.0, 8.6], header_color=ORANGE, font_size=7.9)

# Architecture justification
add_heading(doc, "6. Sustento del planteamiento de arquitectura de solución")
just_rows = [
    ("Adecuación al perfil", "El núcleo agéntico separa comprensión, decisión y ejecución de herramientas. Esto responde al objetivo de detectar el punto de dolor y direccionar al usuario, sin atribuir al agente capacidades que el perfil no autoriza."),
    ("Autonomía controlada", "El LLM propone acciones, pero las derivaciones se validan mediante catálogo, reglas deterministas, confirmación del usuario y registro auditable. El diseño evita improvisar procedimientos o entidades."),
    ("Información vigente", "El conocimiento se administra como corpus actualizable: ingesta con deduplicación, metadatos y recuperación con citas. La respuesta debe mostrar fuentes y vigencia para reducir el riesgo de programas o requisitos desactualizados."),
    ("RAG avanzado", "La combinación de reformulación, búsqueda KNN, deduplicación, reranking y grounding mejora la precisión en un dominio institucional cambiante frente a depender exclusivamente del conocimiento paramétrico del LLM."),
    ("Memoria adecuada", "La memoria corta mantiene continuidad durante la sesión; el resumen de largo plazo conserva antecedentes; la memoria semántica recupera hechos útiles. Separar estas funciones reduce el tamaño del contexto y permite trazabilidad y derecho al olvido."),
    ("Human-in-the-loop", "Los casos de angustia, urgencia o alto impacto económico se convierten en handoff con contexto. La consola en vivo permite que una persona retome la conversación y limita el riesgo de respuestas inadecuadas."),
    ("Escalabilidad y costo", "Cloud Run, Firestore y Storage son servicios administrados y de pago por uso. La separación entre tráfico interactivo y jobs batch permite escalar el canal sin mantener infraestructura ociosa."),
    ("Resiliencia", "El agente es stateless y persiste el contexto. Los fallos de dependencias se manejan con fallback amigable y operaciones asíncronas para no bloquear el webhook; los health checks y alertas facilitan recuperación."),
    ("Seguridad y privacidad", "La firma del canal se valida fail-closed, se aplican guardrails contra inyección y fuga de PII, se restringen permisos con IAM y se gestionan secretos fuera del código. La memoria incorpora borrado/olvido y trazas sanitizadas."),
    ("Medición de calidad", "RAGAS, DeepEval, dataset dorado, trazas y feedback permiten medir fidelidad, relevancia, calidad de la derivación y desempeño de herramientas antes de ampliar autonomía."),
]
add_table(doc, ["Criterio", "Sustento"], just_rows, [4.2, 12.5], header_color=GREEN, font_size=8.15)

add_heading(doc, "7. Controles críticos y criterios de aceptación")
add_bullets(doc, [
    "Derivación: la entidad recomendada debe corresponder a la necesidad y contar con enlace/canal oficial; toda derivación queda registrada con motivo, fuente y fecha.",
    "Vigencia: el contenido con requisitos, plazos, subsidios o programas debe tener propietario, fecha de revisión y proceso de actualización.",
    "Empatía: ante malestar, angustia o temor, la respuesta reconoce la situación, evita promesas y ofrece escalamiento humano cuando corresponda.",
    "Privacidad: no se exponen teléfonos, DNI, RUC u otros identificadores en logs; el acceso al historial se restringe por rol y el usuario puede solicitar eliminación.",
    "Operación: se monitorean latencia, errores, costo por conversación, tasa de handoff, tasa de respuestas con fuente y calidad de recuperación.",
    "Seguridad: se prueban inyección de prompt, poisoning de memoria, abuso de herramientas, contenido externo malicioso y falsificación de webhook.",
])

add_heading(doc, "8. Conclusión")
doc.add_paragraph(
    "La arquitectura propuesta es coherente con SITFEN porque convierte el perfil del agente en una solución operable: canal accesible, núcleo de razonamiento con herramientas, conocimiento oficial recuperable, memoria controlada, derivación restringida y supervisión humana. GCP aporta una base administrada y escalable; LangGraph/LangChain y Gemini permiten implementar la coordinación agéntica; Firestore y Cloud Storage sostienen la memoria y el corpus; y los controles de seguridad, evaluación y observabilidad reducen los riesgos propios de un agente que opera en un contexto de emergencia climática."
)
add_note(doc, "Alcance", "Este documento describe la arquitectura objetivo de la implementación SITFEN v2 y debe complementarse con un inventario de responsables, SLA, costos por ambiente y un plan de actualización del corpus institucional antes del paso a producción.", "FFF5EB", ORANGE)

# Sources / provenance
add_heading(doc, "Anexo A. Fuentes de diseño utilizadas")
add_bullets(doc, [
    "Agent Profile Card – SITFEN v2 (perfil adjunto): dominio, objetivos, conocimiento, herramientas, memoria, autonomía y criticidad.",
    "README.md y CLAUDE.md del repositorio Smart Agentic Chatbot v2: componentes implementados, pipeline del orquestador, despliegue GCP, seguridad y evaluación.",
    "Artefactos gráficos SITFEN del repositorio: arquitectura funcional, despliegue GCP y stack tecnológico.",
])

doc.core_properties.title = "Arquitectura de solución e infraestructura agéntica SITFEN"
doc.core_properties.subject = "Smart Agentic Chatbot v2"
doc.core_properties.author = "OpenAI Codex"
doc.core_properties.keywords = "SITFEN, FEN, arquitectura, agente agéntico, GCP, WhatsApp, RAG"
doc.save(OUT)
print(OUT)
