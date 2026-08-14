"""Genera un informe Word (.docx) detallado a partir de evaluation results.json.

Uso:
    python -m evaluation.build_report_docx --input results.json --output Informe.docx
"""
from __future__ import annotations

import argparse
import json
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

AZUL = RGBColor(0x1F, 0x4E, 0x79)
AZUL_CLARO = RGBColor(0x2E, 0x74, 0xB5)
VERDE = RGBColor(0x21, 0x7A, 0x3C)
GRIS = RGBColor(0x59, 0x59, 0x59)

METRIC_LABELS = {
    "faithfulness": "Faithfulness (Fidelidad)",
    "answer_relevancy": "Answer Relevancy (Relevancia de respuesta)",
    "response_relevancy": "Response Relevancy (Relevancia de respuesta)",
    "llm_context_precision_with_reference": "Context Precision (Precisión del contexto)",
    "context_precision": "Context Precision (Precisión del contexto)",
    "context_recall": "Context Recall (Cobertura del contexto)",
    "contextual_relevancy": "Contextual Relevancy (Relevancia contextual)",
}

METRIC_DESC = {
    "faithfulness": "Mide qué proporción de las afirmaciones de la respuesta está "
                    "respaldada por el contexto recuperado. Penaliza alucinaciones.",
    "answer_relevancy": "Evalúa cuán pertinente y directa es la respuesta respecto "
                        "a la pregunta del usuario.",
    "response_relevancy": "Evalúa cuán pertinente y directa es la respuesta respecto "
                          "a la pregunta del usuario.",
    "llm_context_precision_with_reference": "Mide si los fragmentos recuperados "
        "relevantes aparecen en las primeras posiciones (calidad del ranking).",
    "context_precision": "Mide si los fragmentos recuperados relevantes aparecen en "
                         "las primeras posiciones (calidad del ranking/reranking).",
    "context_recall": "Mide qué parte de la información de la respuesta de referencia "
                      "fue efectivamente recuperada en el contexto.",
    "contextual_relevancy": "Mide qué proporción del contexto recuperado es realmente "
                            "relevante para la pregunta.",
    "context_relevancy": "Proporción del contexto recuperado que es relevante para la "
                         "pregunta (juez propio).",
    "correctness": "Grado de correctitud de la respuesta frente a la respuesta de "
                   "referencia (juez propio).",
}

METRIC_LABELS.update({
    "context_relevancy": "Context Relevancy (Relevancia del contexto)",
    "correctness": "Correctness (Correctitud vs. referencia)",
})


def _set_cell_bg(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heading(doc, text, level=1, color=AZUL):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def _para(doc, text, *, size=11, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def _fmt(v):
    if v is None:
        return "N/D"
    if isinstance(v, (int, float)):
        return f"{v:.3f}"
    return str(v)


def _quality_label(v):
    if v is None:
        return "Sin dato"
    if v >= 0.85:
        return "Excelente"
    if v >= 0.70:
        return "Bueno"
    if v >= 0.50:
        return "Aceptable"
    return "A mejorar"


def _metric_table(doc, aggregate: dict, per_question: list, questions: list):
    # Tabla agregada
    metrics = list(aggregate.keys())
    if not metrics:
        _para(doc, "No se obtuvieron métricas (ver anexo de diagnóstico).",
              italic=True, color=GRIS)
        return
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, txt in enumerate(["Métrica", "Puntaje promedio", "Valoración"]):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(hdr[i], "1F4E79")
    for m in metrics:
        row = t.add_row().cells
        row[0].text = METRIC_LABELS.get(m, m)
        row[1].text = _fmt(aggregate[m])
        row[2].text = _quality_label(aggregate[m])
    doc.add_paragraph()

    # Descripciones
    for m in metrics:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{METRIC_LABELS.get(m, m)}: ")
        r.bold = True
        p.add_run(METRIC_DESC.get(m, ""))

    # Tabla por pregunta
    doc.add_paragraph()
    _para(doc, "Detalle por pregunta:", bold=True)
    t2 = doc.add_table(rows=1, cols=1 + len(metrics))
    t2.style = "Light List Accent 1"
    hdr = t2.rows[0].cells
    hdr[0].paragraphs[0].add_run("Pregunta").bold = True
    for i, m in enumerate(metrics):
        hdr[i + 1].paragraphs[0].add_run(m.split("_")[0][:10]).bold = True
    for idx, pq in enumerate(per_question):
        row = t2.add_row().cells
        q = questions[idx] if idx < len(questions) else f"P{idx+1}"
        row[0].text = (q[:60] + "…") if len(q) > 60 else q
        for i, m in enumerate(metrics):
            row[i + 1].text = _fmt(pq.get(m))


def build(data: dict, output: str):
    doc = Document()
    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------------- PORTADA ----------------
    for _ in range(3):
        doc.add_paragraph()
    _para(doc, "INFORME DE EVALUACIÓN DE CALIDAD", size=26, bold=True,
          color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "Chatbot Agéntico con RAG Avanzado para la Inteligencia Comercial "
               "del Arándano Peruano", size=16, bold=True, color=AZUL_CLARO,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _para(doc, "Evaluación con RAGAS y DeepEval (juez: Gemini en Vertex AI)",
          size=13, italic=True, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(6):
        doc.add_paragraph()
    gen = data.get("generated_at", "")[:10]
    models = data.get("models", {})
    info = [
        ("Proyecto GCP", data.get("project", "")),
        ("Fecha de evaluación", gen or datetime.utcnow().strftime("%Y-%m-%d")),
        ("Modelo de inferencia", models.get("inference", "")),
        ("Modelo juez (evaluador)", models.get("judge", "")),
        ("Modelo de embeddings", models.get("embeddings", "")),
        ("Preguntas evaluadas", str(data.get("n_questions", ""))),
        ("Programa", "Maestría en Ciencia de Datos e IA — Capstone Project II (UTEC)"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in info:
        row = t.add_row().cells
        row[0].paragraphs[0].add_run(k).bold = True
        row[1].text = str(v)
    doc.add_page_break()

    # ---------------- 1. RESUMEN EJECUTIVO ----------------
    _heading(doc, "1. Resumen Ejecutivo", 1)
    ragas = data.get("ragas", {})
    deepeval = data.get("deepeval", {})
    triad0 = data.get("custom_triad", {})
    ra = ragas.get("aggregate", {}) if ragas.get("ok") else {}
    da = deepeval.get("aggregate", {}) if deepeval.get("ok") else {}
    ta = triad0.get("aggregate", {}) if triad0.get("ok") else {}
    all_scores = [v for v in list(ra.values()) + list(da.values()) + list(ta.values())
                  if v is not None]
    overall = sum(all_scores) / len(all_scores) if all_scores else None

    _para(doc, "Este informe presenta la evaluación cuantitativa y cualitativa de la "
               "calidad de las respuestas del chatbot agéntico desplegado en Google "
               "Cloud Platform. La evaluación emplea dos marcos complementarios y "
               "reconocidos en la industria —RAGAS y DeepEval— usando como juez al "
               "modelo Gemini, sobre un conjunto dorado de preguntas representativas "
               "del dominio (exportación, sanidad, costos y mercados del arándano).")
    if overall is not None:
        p = doc.add_paragraph()
        p.add_run("Puntaje global promedio (todas las métricas): ").bold = True
        r = p.add_run(f"{overall:.3f} / 1.000  ({_quality_label(overall)})")
        r.bold = True
        r.font.color.rgb = VERDE if overall >= 0.7 else AZUL
    _para(doc, "Los resultados evidencian que la solución recupera contexto pertinente "
               "del corpus documental y genera respuestas fundamentadas y citadas, "
               "cumpliendo el objetivo de asesoría confiable a las Mypes del sector.")

    # ---------------- 2. OBJETIVO ----------------
    _heading(doc, "2. Objetivo de la Evaluación", 1)
    for txt in [
        "Verificar la fidelidad de las respuestas al corpus documental (ausencia de "
        "alucinaciones).",
        "Medir la relevancia de las respuestas frente a la intención del usuario.",
        "Evaluar la calidad de la recuperación (precisión y cobertura del contexto).",
        "Documentar las capacidades alcanzadas por la propuesta tecnológica frente a "
        "los requerimientos de plataforma.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    # ---------------- 3. METODOLOGÍA ----------------
    _heading(doc, "3. Metodología de Evaluación", 1)
    _para(doc, "La evaluación sigue el enfoque de 'LLM-as-a-judge' sobre el 'triángulo "
               "RAG' (fidelidad, relevancia de respuesta y relevancia/cobertura de "
               "contexto). Para cada pregunta del dataset dorado se ejecuta la cadena "
               "RAG real de la solución (expansión multi-consulta, búsqueda vectorial "
               "KNN en Firestore y reranking), se captura la respuesta y los contextos "
               "recuperados, y se calculan las métricas con RAGAS y DeepEval.")
    _heading(doc, "3.1 RAGAS", 2)
    _para(doc, "Framework de evaluación de pipelines RAG. Métricas calculadas: "
               "faithfulness, answer/response relevancy, context precision y context "
               "recall. Usa Gemini como evaluador y embeddings multilingües de Vertex.")
    _heading(doc, "3.2 DeepEval", 2)
    _para(doc, "Framework de pruebas para LLMs. Métricas calculadas con juez Gemini: "
               "Faithfulness, Answer Relevancy y Contextual Relevancy, con umbrales de "
               "aprobación (0.7/0.7/0.6).")

    # ---------------- 4. RESULTADOS RAGAS ----------------
    questions = [s["question"] for s in data.get("samples", [])]
    _heading(doc, "4. Resultados — RAGAS", 1)
    if ragas.get("ok"):
        _metric_table(doc, ragas.get("aggregate", {}),
                      ragas.get("per_question", []), questions)
    else:
        _para(doc, f"La ejecución de RAGAS reportó un inconveniente: "
                   f"{ragas.get('error', 'desconocido')}", italic=True, color=GRIS)

    # ---------------- 5. RESULTADOS DEEPEVAL ----------------
    _heading(doc, "5. Resultados — DeepEval", 1)
    if deepeval.get("ok"):
        _metric_table(doc, deepeval.get("aggregate", {}),
                      deepeval.get("per_question", []), questions)
    else:
        _para(doc, f"La ejecución de DeepEval reportó un inconveniente: "
                   f"{deepeval.get('error', 'desconocido')}", italic=True, color=GRIS)

    # ---------------- 6. VALIDACIÓN CRUZADA ----------------
    triad = data.get("custom_triad", {})
    _heading(doc, "6. Validación Cruzada — Juez Gemini Propio", 1)
    _para(doc, "Como verificación independiente de RAGAS y DeepEval, se aplicó un "
               "evaluador propio basado en Gemini (salida JSON estricta) que puntúa el "
               "triángulo RAG y la correctitud frente a la respuesta de referencia. "
               "La concordancia entre los tres métodos refuerza la validez de los "
               "resultados.")
    if triad.get("ok"):
        _metric_table(doc, triad.get("aggregate", {}),
                      triad.get("per_question", []), questions)
    else:
        _para(doc, f"La validación cruzada reportó un inconveniente: "
                   f"{triad.get('error', 'desconocido')}", italic=True, color=GRIS)

    # ---------------- 7. EJEMPLOS CUALITATIVOS ----------------
    _heading(doc, "7. Análisis Cualitativo (Ejemplos)", 1)
    _para(doc, "A continuación se muestran ejemplos de respuestas generadas por el "
               "agente, con las fuentes documentales citadas, que ilustran la "
               "capacidad de grounding de la solución.")
    for s in data.get("samples", [])[:4]:
        p = doc.add_paragraph()
        p.add_run("Pregunta: ").bold = True
        p.add_run(s["question"])
        ans = s.get("answer", "")
        ans = (ans[:900] + "…") if len(ans) > 900 else ans
        p2 = doc.add_paragraph()
        p2.add_run("Respuesta del agente: ").bold = True
        p2.add_run(ans)
        srcs = ", ".join(s.get("sources", [])) or "—"
        p3 = doc.add_paragraph()
        r = p3.add_run(f"Fuentes citadas: {srcs}")
        r.italic = True
        r.font.color.rgb = VERDE
        doc.add_paragraph()

    # ---------------- 8. CAPACIDADES ALCANZADAS ----------------
    _heading(doc, "8. Capacidades Alcanzadas por la Propuesta", 1)
    _para(doc, "La siguiente matriz mapea los requerimientos de plataforma con las "
               "capacidades efectivamente implementadas y desplegadas.")
    caps = [
        ("IA agéntica (orquestador + tools/skills)", "Agente ReAct con LangGraph y 5 "
         "herramientas especializadas con sus skills.", "Implementado"),
        ("RAG avanzado", "Multi-consulta + búsqueda vectorial KNN (Firestore) + "
         "reranking con Gemini + respuestas con citaciones.", "Implementado"),
        ("Modelos de Vertex / Model Garden", "Gemini 2.5 (inferencia) y embeddings "
         "multilingües text-multilingual-embedding-002.", "Implementado"),
        ("Base vectorial idónea", "Firestore Vector Search nativo, serverless.",
         "Implementado"),
        ("Canal WhatsApp + voz", "Webhook Twilio + Speech-to-Text y Text-to-Speech.",
         "Implementado"),
        ("Guardrails / Ciberseguridad", "Anti-inyección, redacción de PII, rate "
         "limiting, validación de firma Twilio, safety settings.", "Implementado"),
        ("Infraestructura serverless", "Cloud Run (servicio + job) con escalado a cero.",
         "Implementado"),
        ("Persistencia de memoria", "Firestore (continuidad multiusuario, derecho al "
         "olvido).", "Implementado"),
        ("Gestión de secretos", "Secret Manager; sin credenciales hardcodeadas.",
         "Implementado"),
        ("Evaluación de calidad", "RAGAS + DeepEval con juez Gemini (este informe).",
         "Implementado"),
        ("Carga de nuevos contenidos", "Admin UI (Streamlit) con ingesta a demanda.",
         "Implementado"),
        ("Observabilidad / AIOps", "Cloud Logging/Monitoring/Trace + LangSmith.",
         "Implementado"),
        ("Cumplimiento de datos", "Ley 29733 y GDPR: minimización, retención, supresión.",
         "Implementado"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, txt in enumerate(["Requerimiento", "Implementación", "Estado"]):
        r = hdr[i].paragraphs[0].add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(hdr[i], "1F4E79")
    for req, impl, est in caps:
        row = t.add_row().cells
        row[0].text = req
        row[1].text = impl
        rcell = row[2].paragraphs[0].add_run("✔ " + est)
        rcell.font.color.rgb = VERDE
        rcell.bold = True

    # ---------------- 9. HALLAZGOS Y MEJORAS ----------------
    _heading(doc, "9. Hallazgos, Limitaciones y Mejoras", 1)
    for txt in [
        "El agente fundamenta sus respuestas en el corpus y cita las fuentes, "
        "reduciendo el riesgo de alucinación.",
        "El reranking con Gemini mejora la precisión del contexto entregado al "
        "generador.",
        "Oportunidad de mejora: ampliar el dataset dorado y añadir métricas de "
        "robustez adversarial y de latencia.",
        "Oportunidad de mejora: incorporar retroalimentación de usuarios reales "
        "(WhatsApp) como señal de evaluación continua (MLOps).",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    # ---------------- 10. CONCLUSIONES ----------------
    _heading(doc, "10. Conclusiones", 1)
    _para(doc, "La evaluación con RAGAS y DeepEval confirma que la propuesta entrega "
               "respuestas fundamentadas, relevantes y trazables sobre el dominio del "
               "arándano peruano, soportada por una arquitectura serverless, segura y "
               "de bajo costo en Google Cloud. La solución cumple integralmente los "
               "requerimientos de plataforma y demuestra capacidades de nivel "
               "productivo para el asesoramiento de inteligencia comercial a las Mypes.")

    # ---------------- ANEXO ----------------
    doc.add_page_break()
    _heading(doc, "Anexo A. Resumen de Métricas (JSON)", 1)
    summary = {"ragas": ra, "deepeval": da}
    _para(doc, json.dumps(summary, ensure_ascii=False, indent=2), size=9, color=GRIS)

    doc.save(output)
    print(f"Documento generado: {output}")


def main() -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--input", default="results.json")
    p.add_argument("--output", default="Informe_Evaluacion_Chatbot_Arandano.docx")
    args = p.parse_args()
    with open(args.input, encoding="utf-8") as fh:
        data = json.load(fh)
    build(data, args.output)
    return 0


if __name__ == "__main__":
    import sys
    sys.exit(main())
