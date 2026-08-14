"""Carga de documentos PDF desde disco local o Google Cloud Storage.

Extrae texto por página con pdfplumber (mejor manejo de tablas/columnas que
pypdf para los reportes y tarifarios del corpus) y cae a pypdf si falla.
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterator, Optional

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    source: str  # nombre de archivo
    title: str
    pages: list[str] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(self.pages)


def _clean(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pages(raw: bytes) -> list[str]:
    pages: list[str] = []
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                pages.append(_clean(page.extract_text() or ""))
        if any(pages):
            return pages
    except Exception as exc:  # noqa: BLE001
        logger.warning("pdfplumber falló, usando pypdf: %s", exc)

    pages = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        for page in reader.pages:
            pages.append(_clean(page.extract_text() or ""))
    except Exception as exc:  # noqa: BLE001
        logger.error("No se pudo extraer texto del PDF: %s", exc)
    return pages


def _extract_docx(raw: bytes) -> list[str]:
    """Extrae texto de un .docx (párrafos + celdas de tablas)."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = _clean("\n".join(parts))
        return [text] if text else []
    except Exception as exc:  # noqa: BLE001
        logger.error("No se pudo extraer texto del DOCX: %s", exc)
        return []


def _extract_txt(raw: bytes) -> list[str]:
    """Extrae texto de un .txt (utf-8 con respaldo latin-1)."""
    for enc in ("utf-8", "latin-1"):
        try:
            text = _clean(raw.decode(enc))
            return [text] if text else []
        except Exception:  # noqa: BLE001
            continue
    return []


# Extractor por extensión.
_EXTRACTORS = {".pdf": _extract_pages, ".docx": _extract_docx, ".txt": _extract_txt}
SUPPORTED_EXTS = tuple(_EXTRACTORS.keys())


def _extract(name: str, raw: bytes) -> list[str]:
    ext = Path(name).suffix.lower()
    extractor = _EXTRACTORS.get(ext)
    if not extractor:
        logger.warning("Formato no soportado (%s): %s", ext, name)
        return []
    return extractor(raw)


def _make_document(name: str, raw: bytes) -> Optional[LoadedDocument]:
    # Validación de integridad + seguridad ANTES de extraer/indexar (lanza
    # ValidationError si el archivo es inválido, corrupto o sospechoso).
    from ingestion.validation import validate_file

    validate_file(name, raw)
    pages = _extract(name, raw)
    if not any(pages):
        return None
    base = Path(name).name
    return LoadedDocument(
        source=base,
        title=Path(base).stem.replace("-", " ").replace("_", " ").strip(),
        pages=pages,
    )


def load_local_documents(folder: str | Path) -> Iterator[LoadedDocument]:
    """Carga PDF/DOCX/TXT desde una carpeta local, recorriendo subcarpetas.

    `source` es siempre el nombre base del archivo (sin ruta de carpeta), por
    lo que debe mantenerse único en todo el árbol de `folder` para evitar
    colisiones entre subcarpetas distintas.
    """
    from ingestion.validation import ValidationError

    folder = Path(folder)
    for path in sorted(folder.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTS:
            continue
        try:
            doc = _make_document(path.name, path.read_bytes())
        except ValidationError as exc:
            logger.warning("Validación falló, se omite %s: %s", path.name, exc)
            continue
        if doc is None:
            logger.warning("Sin texto extraíble: %s", path.name)
            continue
        yield doc


def load_gcs_documents(bucket_name: str, prefix: str = "") -> Iterator[LoadedDocument]:
    """Carga PDF/DOCX/TXT desde un bucket de GCS."""
    from google.cloud import storage

    from ingestion.validation import ValidationError

    client = storage.Client()
    bucket = client.bucket(bucket_name)
    for blob in client.list_blobs(bucket, prefix=prefix):
        if not blob.name.lower().endswith(SUPPORTED_EXTS):
            continue
        try:
            doc = _make_document(blob.name, blob.download_as_bytes())
        except ValidationError as exc:
            logger.warning("Validación falló, se omite %s: %s", blob.name, exc)
            continue
        if doc is None:
            logger.warning("Sin texto extraíble en GCS: %s", blob.name)
            continue
        yield doc


def load_single_gcs_document(bucket_name: str, blob_name: str) -> Optional[LoadedDocument]:
    """Carga un único PDF/DOCX/TXT de GCS."""
    from google.cloud import storage

    client = storage.Client()
    blob = client.bucket(bucket_name).blob(blob_name)
    return _make_document(blob_name, blob.download_as_bytes())


# --------------------------- Páginas web (scraping) ---------------------------
class _TextExtractor(HTMLParser):
    """Extrae texto visible omitiendo navegación, scripts y estilos."""

    _SKIP = {"script", "style", "noscript", "head", "nav", "footer",
             "header", "aside", "form", "svg", "button"}

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._skip = 0

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip += 1

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skip > 0:
            self._skip -= 1

    def handle_data(self, data):
        if self._skip == 0 and data.strip():
            self.parts.append(data.strip())


def _html_title(html: str) -> str:
    m = re.search(r"<title[^>]*>(.*?)</title>", html, re.I | re.S)
    return _clean(re.sub(r"\s+", " ", m.group(1))) if m else ""


def _html_to_text(html: str) -> str:
    """Extrae el contenido principal: trafilatura si está, si no stdlib."""
    try:
        import trafilatura

        extracted = trafilatura.extract(
            html, include_comments=False, include_tables=True, favor_recall=True
        )
        if extracted and extracted.strip():
            return _clean(extracted)
    except Exception as exc:  # noqa: BLE001
        logger.info("trafilatura no disponible/falló, usando extractor básico: %s", exc)

    parser = _TextExtractor()
    try:
        parser.feed(html)
    except Exception:  # noqa: BLE001
        pass
    return _clean("\n".join(parser.parts))


def load_url(url: str, *, timeout: float = 20.0) -> Optional[LoadedDocument]:
    """Descarga una página web preseleccionada y la convierte en documento.

    La 'source' es la propia URL, por lo que aparece y se gestiona en el listado
    igual que los archivos (chunks, re-ingesta/replace, borrado).
    """
    import httpx

    headers = {"User-Agent": "Mozilla/5.0 (compatible; AruBot/1.0; corpus-arandano)"}
    try:
        resp = httpx.get(url, headers=headers, timeout=timeout, follow_redirects=True)
        resp.raise_for_status()
    except Exception as exc:  # noqa: BLE001
        logger.error("No se pudo descargar la URL %s: %s", url, exc)
        return None

    ctype = resp.headers.get("content-type", "").lower()
    if "application/pdf" in ctype or url.lower().endswith(".pdf"):
        pages = _extract_pages(resp.content)
        title = Path(url).stem.replace("-", " ").replace("_", " ").strip()
    else:
        text = _html_to_text(resp.text)
        pages = [text] if text else []
        title = _html_title(resp.text) or url
    if not any(pages):
        logger.warning("Sin texto extraíble de la URL: %s", url)
        return None
    return LoadedDocument(source=url, title=title, pages=pages)


def load_local_pdfs(folder: str | Path) -> Iterator[LoadedDocument]:
    folder = Path(folder)
    for pdf_path in sorted(folder.glob("*.pdf")):
        raw = pdf_path.read_bytes()
        pages = _extract_pages(raw)
        if not any(pages):
            logger.warning("Sin texto extraíble: %s", pdf_path.name)
            continue
        yield LoadedDocument(
            source=pdf_path.name,
            title=pdf_path.stem.replace("-", " ").replace("_", " ").strip(),
            pages=pages,
        )


def load_gcs_pdfs(bucket_name: str, prefix: str = "") -> Iterator[LoadedDocument]:
    from google.cloud import storage

    client = storage.Client()
    bucket = client.bucket(bucket_name)
    for blob in client.list_blobs(bucket, prefix=prefix):
        if not blob.name.lower().endswith(".pdf"):
            continue
        raw = blob.download_as_bytes()
        pages = _extract_pages(raw)
        if not any(pages):
            logger.warning("Sin texto extraíble en GCS: %s", blob.name)
            continue
        name = Path(blob.name).name
        yield LoadedDocument(
            source=name,
            title=Path(name).stem.replace("-", " ").replace("_", " ").strip(),
            pages=pages,
        )


def load_single_gcs_pdf(bucket_name: str, blob_name: str) -> Optional[LoadedDocument]:
    from google.cloud import storage

    client = storage.Client()
    blob = client.bucket(bucket_name).blob(blob_name)
    raw = blob.download_as_bytes()
    pages = _extract_pages(raw)
    if not any(pages):
        return None
    name = Path(blob_name).name
    return LoadedDocument(
        source=name,
        title=Path(name).stem.replace("-", " ").replace("_", " ").strip(),
        pages=pages,
    )
